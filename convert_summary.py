from docx import Document

# Read markdown file
with open("summary.md", "r") as file:
    content = file.read()

# Create docx
doc = Document()
doc.add_heading("Employee Attrition Prediction - Summary", 0)

for line in content.split("\n"):
    doc.add_paragraph(line)

doc.save("summary.docx")

print("summary.docx created successfully!")
